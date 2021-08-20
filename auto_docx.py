<<<<<<< HEAD
#author:whale3070
#version:v1
#time:2021-07-21
import configparser
import os
import re
from docx import Document
from docx.shared import Inches


def Use_template(Select_Num):
    #获取当前目录下的所有文件
    path = os.listdir(os.getcwd())
    company_list = company()
    # print(company_list)
    '''
    ['重庆赣江实业有限公司', '新希望六和股份有限公司', '白馬窯業']
    '''
    # 漏洞地址名称
    vul_address_list = vul_address()
    # 使用默认模板时
    if Select_Num == 1:
        # 读取当前目录下文件

        if "CNVD.ini" in path:
            # 读取配置文件
            config = configparser.ConfigParser()
            config.read("CNVD.ini", encoding="utf-8")
            # 获取所有的选项
            secs = config.sections()
            # print(secs)
            i = 0
            # 遍历所有的模板名称并可选择
            for section in secs:
                print("编号" + str(i) + ":" + section)
                i += 1
            # 锁定模板文件
            sec_num = int(input("请输入对应模板文件的编号:"))
            #漏洞名称
            title_texts = config.get(secs[sec_num], "title_texts")
            #漏洞发现
            Vuln_discovery = config.get(secs[sec_num], "Vuln_discovery")
            #漏洞地址
            Vuln_address = config.get(secs[sec_num], "Vuln_address")
            #漏洞验证
            Vuln_check = config.get(secs[sec_num], "Vuln_check")
            new_docx(company_list, vul_address_list, Vuln_discovery, Vuln_address, Vuln_check, title_texts)
        else:
            print("读取配置文件失败，请检查CNVD.ini文件是否存在，是否生成新的自定义配置选项,使用请按1，退出请按0:")
            flag = int(input())
            if flag:
                Create_New_temp()
            else:
                exit(0)
    # 创建新的模板文件
    else:

        Create_New_temp()

def Create_New_temp():
    # 自定义信息，并询问是否需要生成到模板中
    title_texts = input("请输入漏洞名称:")
    Vuln_discovery = input("请输入漏洞发现内容:")
    Vuln_address = input("请输入漏洞地址:")
    Vuln_check = input("请输入漏洞证明:")

    template_name = input("请输入配置名称(请勿输入已存在配置名称):")
    conf = configparser.ConfigParser()
    conf.add_section(template_name)
    conf.set(template_name, "title_texts", title_texts)
    conf.set(template_name, "Vuln_discovery", Vuln_discovery)
    conf.set(template_name, "Vuln_address", Vuln_address)
    conf.set(template_name, "Vuln_check", Vuln_check)

    conf.write(open("CNVD.ini", 'a', encoding="utf-8"))
    print("是否生成docx文档,生成请输入1，退出请按0")
    num = int(input())
    if num ==1:
        Use_template(1)
    else:
        exit(0)

# 获取公司名称
def company():
    list = []
    company_file = open("company.txt", "r", encoding='utf-8')
    for eachCompany in company_file.readlines():
        list.append(eachCompany.strip("\n"))
    company_file.close()
    return list
 #获取漏洞地址
def vul_address():
    list = []
    vul_add = open("pwned.txt", "r", encoding='utf-8')
    for each in vul_add.readlines():
        mysplit = each.strip("\n").split(",")
        #print(mysplit)
        list.append(mysplit)
    vul_add.close()
    return list

# def get_file_path_by_name(file_dir):
#     '''
#     获取指定路径下所有文件的绝对路径
#     :param file_dir:
#     :return:
#     '''
#     L = []
#     for root,dirs,files in os.walk(file_dir):  # 获取所有文件
#         for file in files:  # 遍历所有文件名
#             if os.path.splitext(file)[1] == '.bmp':
#                 L.append(os.path.join(root, file))  # 拼接处绝对路径并放入列表
#     print('总文件数目：', len(L))
#     return L

    
def mk_docx(num,company_name,vul_address_list, Vuln_discovery, Vuln_address, Vuln_check, title_texts):
    #print(num)
    picnum=0
    for vul_address_name in vul_address_list:
        #trueIp =re.search(r'(([01]{0,1}\d{0,1}\d|2[0-4]\d|25[0-5])\.){3}([01]{0,1}\d{0,1}\d|2[0-4]\d|25[0-5])',vul_address_name)
        # trueIp = re.findall( r'[0-9]+(?:\.[0-9]+){3}',vul_address_name)
        # print(trueIp)
        print(vul_address_name)
        '''
        pic_path = ''
        for i in img_path_list:
            #print(i)
            img_path = re.findall(trueIp[0],i)
            if len(img_path)>0:
                #print(i)
                
                pic_path = i
        print(pic_path)
        '''
        document = Document()
        document.add_heading(company_name+'存在'+title_texts, 0)
        #document.add_picture(pic_path,width=Inches(6.0) )

        document.add_heading('漏洞发现', level=1)
        p = document.add_paragraph(Vuln_discovery)
        document.add_heading('漏洞地址', level=1)
        Vuln_address = vul_address_name+Vuln_address
        p1 = document.add_paragraph(Vuln_address)

        document.add_heading('漏洞证明', level=1)
        p2 = document.add_paragraph(Vuln_check)
        documentName = company_name+'.docx'
        
        
        #print('title: '+company_name.strip('\n')+'存在蓝凌OA custom.jsp 任意文件读取漏洞')
        picnum+=1
        print(str(num)+'-'+str(picnum)+": "+'vul_address: '+vul_address_name.strip('\n'))
        document.save(str(num)+'-'+str(picnum)+documentName)
        
        

        
def new_docx(company_list,vul_address_list, Vuln_discovery, Vuln_address, Vuln_check, title_texts):
    num=0

    for j in range(len(vul_address_list)):
        num+=1
        mk_docx(num,company_list[j],vul_address_list[j], Vuln_discovery, Vuln_address, Vuln_check, title_texts)

def auto_docx():
    #图片读取写入,暂时不写.
    # img_path_list = get_file_path_by_name('C:\\Users\\Local\\Desktop\\new\\images')
    print("欢迎使用CNVD模板生成工具，是否使用系统默认模板文件,使用请按1，否则请按0:")
    Select_Num = int(input())
    # 启动生成模板文件
    Use_template(Select_Num)








    # company_list = company()
    # #print(company_list)
    # '''
    # ['重庆赣江实业有限公司', '新希望六和股份有限公司', '白馬窯業']
    # '''
    # #漏洞地址名称
    # vul_address_list = vul_address()
    # #print(vul_address_list[1])
    # '''
    # ['http://61.163.234.38:8090/sys/ui/extend/varkind/custom.jsp', 'http://119.249.96.46:8090/sys/ui/extend/varkind/custom.jsp', 'http://116.177.226.119:8090/sys/ui/extend/varkind/custom.jsp', 'http://221.193.246.90:8090/sys/ui/extend/varkind/custom.jsp', 'http://117.157.243.34:8090/sys/ui/extend/varkind/custom.jsp', 'http://116.211.78.155:8090/sys/ui/extend/varkind/custom.jsp', 'http://112.90.245.13:8090/sys/ui/extend/varkind/custom.jsp', 'http://117.187.245.56:8090/sys/ui/extend/varkind/custom.jsp', 'http://111.206.10.146:8090/sys/ui/extend/varkind/custom.jsp', 'http://14.205.40.192:8090/sys/ui/extend/varkind/custom.jsp', 'http://111.43.179.221:8090/sys/ui/extend/varkind/custom.jsp', 'http://113.201.107.205:8090/sys/ui/extend/varkind/custom.jsp']
    # '''

if __name__ == '__main__':
    auto_docx()
=======
#coding: utf-8
# author:whale3070，muyu
# version:v3
# time:2021-08-19

import configparser
import os
from docx import Document
import time
import locale
import sys

locale.setlocale(locale.LC_CTYPE, 'Chinese')
locale.setlocale(locale.LC_ALL,'en')


def Use_template(Select_Num):
    # 获取当前目录下的所有文件
    path = os.listdir(os.getcwd())
    company_list = company()
    # print(company_list)
    # 漏洞地址名称
    vul_address_list = vul_address()
    # print(len(company_list))
    # print(len(vul_address_list))
    if len(company_list) != len(vul_address_list):
        print("公司列表与对应url列表信息不一致，请核对数据！！")
        sys.exit(0)
    # 使用默认模板时
    if Select_Num == 1:
        # 读取当前目录下文件

        if "CNVD.ini" in path:
            # 读取配置文件

            config = configparser.ConfigParser()
            config.read("CNVD.ini", encoding="utf-8")

            # 获取所有的选项

            secs = config.sections()
            if not secs:
                print("配置文件为空或格式错误，请检查配置文件是否正确，程序已退出！")
                sys.exit(0)
            # print(secs)
            i = 0
            # 遍历所有的模板名称并可选择
            for section in secs:
                print("编号" + str(i) + ":" + section)
                i += 1
            # 锁定模板文件
            sec_num = int(input("请输入对应模板文件的编号:"))
            title_texts = config.get(secs[sec_num], "title_texts")
            Vuln_type_texts = config.get(secs[sec_num], "Vuln_type_texts")
            Vuln_hazard_texts = config.get(secs[sec_num], "Vuln_hazard_texts")
            Vuln_intrude = config.get(secs[sec_num], "Vuln_intrude")
            Submitter = config.get(secs[sec_num], "Submitter")
            mail_summitter = config.get(secs[sec_num], "mail_summitter")
            vuln_local = config.get(secs[sec_num], "vuln_local")
            vuln_need_login = int(config.get(secs[sec_num], "vuln_need_login"))

            # 漏洞详情
            vuln_check = config.get(secs[sec_num], "vuln_check")

            new_docx(company_list, vul_address_list, title_texts, Vuln_type_texts, Vuln_hazard_texts, Vuln_intrude,
                     Submitter, mail_summitter, vuln_local,
                     vuln_need_login, vuln_check)

            # print("读取配置文件失败，请检查配置文件是否正常！")
        else:
            print("读取配置文件失败，请检查CNVD.ini文件是否存在，是否生成新的自定义配置选项,使用请按1，退出请按0:")
            flag = int(input())
            if flag:
                Create_New_temp()
            else:
                exit(0)
    # 创建新的模板文件
    else:

        Create_New_temp()


def Create_New_temp():
    # 自定义信息，并询问是否需要生成到模板中
    title_texts = input("请输入漏洞名称:")
    Vuln_type_texts = input("请输入漏洞类型:")
    Vuln_hazard_texts = input("请输入漏洞危害:")
    Vuln_intrude = input("请输入漏洞简介:")
    Submitter = input("请输入提交人员信息:")
    mail_summitter = input("请输入提交人员邮箱:")
    vuln_local = input("请输入漏洞定位URL:")
    vuln_need_login = input("请输入漏洞触发条件，1为需要登陆，0为不需要登陆:")
    vuln_check = input("请输入漏洞详情信息:")
    vuln_check = vuln_check.replace("%", "%%")
    template_name = input("请输入配置名称(请勿输入已存在配置名称):")

    conf = configparser.ConfigParser()
    conf.add_section(template_name)
    conf.set(template_name, "title_texts", title_texts)
    conf.set(template_name, "Vuln_type_texts", Vuln_type_texts)
    conf.set(template_name, "Vuln_hazard_texts", Vuln_hazard_texts)
    conf.set(template_name, "Vuln_intrude", Vuln_intrude)
    conf.set(template_name, "Submitter", Submitter)
    conf.set(template_name, "mail_summitter", mail_summitter)
    conf.set(template_name, "vuln_local", vuln_local)
    conf.set(template_name, "vuln_need_login", vuln_need_login)
    conf.set(template_name, "vuln_check", vuln_check)

    conf.write(open("CNVD.ini", 'a', encoding="utf-8"))
    print("是否生成docx文档,生成请输入1，退出请按0")
    num = int(input())
    if num == 1:
        Use_template(1)
    else:
        sys.exit(0)


# 获取公司名称
def company():
    list = []
    company_file = open("company.txt", "r", encoding='utf-8')
    for eachCompany in company_file.readlines():
        list.append(eachCompany.strip("\n"))
    company_file.close()
    return list


# 获取漏洞地址
def vul_address():
    list = []
    vul_add = open("pwned.txt", "r", encoding='utf-8')
    for each in vul_add.readlines():
        mysplit = each.strip("\n")
        # print(mysplit)
        list.append(mysplit)
    vul_add.close()
    return list


def mk_docx(company_name, vul_address_list, title_texts, Vuln_type_texts, Vuln_hazard_texts, Vuln_intrude_texts,
            Submitter_texts, mail_summitter_texts, vuln_local_texts,
            vuln_need_login, vuln_check,num):
    # print(vul_address_list)
    asb_path = os.getcwd()
    all_path = asb_path + "\\test.docx"
    document = Document(all_path)
    # print(template_num)

    tables = document.tables
    # 定位输出点
    for table in tables:

        # 定位漏洞名称
        title = table.cell(1, 1)
        title_texts = company_name + "存在" + title_texts
        title.text = title_texts


        # 定位漏洞类型
        Vuln_type = table.cell(3, 1)
        # vuln_type_text =  Vuln_type_texts
        Vuln_type.text = Vuln_type_texts
        # print(Vuln_type.text)

        # 漏洞危害登记
        Vuln_hazard = table.cell(4, 1)
        # Vuln_hazard_text = Vuln_hazard_texts
        Vuln_hazard.text = Vuln_hazard_texts
        # print(Vuln_hazard.text)

        # 漏洞简介
        Vuln_intrude = table.cell(5, 1)
        # Vuln_intrude_text = Vuln_intrude_texts
        Vuln_intrude.text = Vuln_intrude_texts
        # print(Vuln_intrude.text)

        # 定位人员信息
        Submitter = table.cell(8, 1)
        # Submitter_text = Submitter_texts
        Submitter.text = Submitter_texts
        # print(Submitter.text)

        # 定位提交邮箱
        mail_summitter = table.cell(9, 1)
        # mail_summitter_text = mail_summitter_texts
        mail_summitter.text = mail_summitter_texts
        # print(mail_summitter.text)

        # 定位提交时间
        date_submit = table.cell(10, 1)
        date_submit_text = time.strftime("%Y年%m月%d日".encode('unicode_escape').decode('utf8'), time.localtime()).encode('utf-8').decode('unicode_escape')
        date_submit.text = date_submit_text
        # print(date_submit.text)

        # 漏洞定位
        vuln_local = table.cell(12, 1)
        # 田哥定制版
        # vuln_url = vul_address_list + vuln_local_texts
        vuln_local_text = "存在漏洞的界面:" + vul_address_list + vuln_local_texts
        vuln_local.text = vuln_local_text
        # print(vuln_local.text)

        # 漏洞触发条件
        select_one = "无，无需登录"
        select_two = "需要登陆"
        if vuln_need_login:
            table.cell(13, 1).text = select_two
        else:
            table.cell(13, 1).text = select_one
        # print(table.cell(13,1).text)

        # 漏洞验证过程
        # 将漏洞详情置空，用户填写
        table_content = table.cell(14, 1)
        # table_content.text = vuln_check + "\n " + vuln_url
        table_content.text = vuln_check.replace("%%","%")

        save_path = asb_path + "\\result\\" + time.strftime("%Y-%m-%d", time.localtime())
        if not os.path.exists(save_path):
            os.makedirs(save_path)

        # print("C:\\Users\\Administrator\\Desktop\\[事件型漏洞" +str(template_num_flag)+ title_texts + "].docx")

        document.save(save_path + "\\[事件型漏洞" +str(num)+"]"+title_texts + ".docx")


def new_docx(company_list, vul_address_list, title_texts, Vuln_type_texts, Vuln_hazard_texts, Vuln_intrude, Submitter,
             mail_summitter, vuln_local,
             vuln_need_login, vuln_check):
    num = 0
    for j in range(len(vul_address_list)):
        num+=1
        mk_docx(company_list[j], vul_address_list[j], title_texts, Vuln_type_texts, Vuln_hazard_texts, Vuln_intrude,
                Submitter, mail_summitter, vuln_local,
                vuln_need_login, vuln_check,num)


def auto_docx():
    # 图片读取写入,暂时不写.
    # img_path_list = get_file_path_by_name('C:\\Users\\Local\\Desktop\\new\\images')
    print("欢迎使用CNVD模板生成工具，是否使用系统默认模板文件,使用请按1，否则请按0:")
    Select_Num = int(input())
    # 启动生成模板文件
    Use_template(Select_Num)


if __name__ == '__main__':
    auto_docx()
>>>>>>> 8ebfeca (V3.0)
